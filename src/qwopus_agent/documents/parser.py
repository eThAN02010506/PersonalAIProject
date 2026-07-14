"""Document-to-Markdown parser.

All unstructured documents are normalized to Markdown before later indexing or LLM analysis. This
keeps downstream modules independent from PDF/DOCX/TXT format details.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from qwopus_agent.documents.mineru import MinerUUnavailableError, parse_document_with_mineru


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


@dataclass(frozen=True)
class ParsedDocument:
    """Parsed Markdown representation of one uploaded document."""

    # 原因：报告和 UI 需要知道文档来自哪里。
    # 作用：保存源文件路径。
    source_path: Path

    # 原因：所有后续分析统一处理 Markdown。
    # 作用：保存解析后的 Markdown 文本。
    markdown: str

    # 原因：UI 需要展示简要元数据，测试也需要可验证的结构。
    # 作用：保存页数、段落数、字符数等轻量信息。
    metadata: dict[str, int | str]


def parse_document(file_path: str | Path) -> ParsedDocument:
    """Parse PDF, DOCX, Markdown, or TXT into Markdown."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix or '<none>'}")

    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".md":
        markdown = path.read_text(encoding="utf-8", errors="ignore")
        return _build_parsed_document(path, markdown, source_type="markdown")

    text = path.read_text(encoding="utf-8", errors="ignore")
    return _build_parsed_document(path, text, source_type="text")


def _parse_pdf(path: Path) -> ParsedDocument:
    """Extract text from PDF pages."""
    try:
        return _parse_with_mineru(path, source_type="pdf")
    except MinerUUnavailableError:
        pass

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            # 原因：保留页码能让后续分析结果可追溯。
            # 作用：把 PDF 页内容转换成 Markdown 小节。
            pages.append(f"## Page {index}\n\n{text.strip()}")

    markdown = "\n\n".join(pages)
    parsed = _build_parsed_document(path, markdown, source_type="pdf")
    parsed.metadata["parser"] = "pypdf"
    parsed.metadata["pages"] = len(reader.pages)
    return parsed


def _parse_docx(path: Path) -> ParsedDocument:
    """Extract paragraphs and tables from DOCX into Markdown."""
    try:
        return _parse_with_mineru(path, source_type="docx")
    except MinerUUnavailableError:
        pass

    from docx import Document

    document = Document(str(path))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            # 原因：DOCX 表格需要保留结构，便于后续分析和展示。
            # 作用：把表格转换成 Markdown table。
            blocks.append(f"### Table {table_index}\n\n{_rows_to_markdown_table(rows)}")

    return _build_parsed_document(path, "\n\n".join(blocks), source_type="docx")


def _parse_with_mineru(path: Path, source_type: str) -> ParsedDocument:
    mineru_result = parse_document_with_mineru(path)
    parsed = _build_parsed_document(path, mineru_result.markdown, source_type=source_type)
    # 原因：调试和后续报告需要知道文档实际由哪个解析器处理。
    # 作用：标记 MinerU 输出路径，便于排查解析质量。
    parsed.metadata["parser"] = "mineru"
    parsed.metadata["mineru_command"] = mineru_result.command
    parsed.metadata["mineru_output_path"] = str(mineru_result.output_path)
    return parsed


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    body = normalized[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _build_parsed_document(path: Path, markdown: str, source_type: str) -> ParsedDocument:
    lines = [line for line in markdown.splitlines() if line.strip()]
    return ParsedDocument(
        source_path=path,
        markdown=markdown,
        metadata={
            "source_type": source_type,
            "characters": len(markdown),
            "non_empty_lines": len(lines),
            "words": len(markdown.split()),
        },
    )
