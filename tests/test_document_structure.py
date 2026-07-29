import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from qwopus_agent.documents import (
    DocumentStore,
    build_document_structure,
    chunk_document_structure,
    parse_document,
    summarize_document,
)
from qwopus_agent.documents.mineru import MinerUUnavailableError


class DocumentStructureTests(unittest.TestCase):
    def test_markdown_headings_build_parent_paths_without_using_pages_as_sections(self) -> None:
        structure = build_document_structure(
            "# Chapter One\n\nIntro.\n\n## Page 2\n\n## Risks\n\n"
            "Risk details.\n\n### Suppliers\n\nSupplier details.",
            source="report.pdf",
        )

        self.assertEqual(
            [section.section_path for section in structure.sections],
            [
                ("Chapter One",),
                ("Chapter One", "Risks"),
                ("Chapter One", "Risks", "Suppliers"),
            ],
        )
        self.assertEqual(structure.sections[1].page_start, 2)
        self.assertEqual(structure.sections[2].parent_id, structure.sections[1].id)

    def test_plain_text_numbered_headings_are_inferred(self) -> None:
        structure = build_document_structure(
            "第一章 背景\n背景正文\n1.1 目标\n目标正文\n1.2 范围\n范围正文",
            source="requirements.txt",
        )

        self.assertEqual(structure.sections[0].title, "第一章 背景")
        self.assertEqual(structure.sections[1].section_path, ("第一章 背景", "1.1 目标"))
        self.assertEqual(structure.sections[2].section_path, ("第一章 背景", "1.2 范围"))

    def test_docx_mineru_bold_numbered_headings_are_inferred(self) -> None:
        structure = build_document_structure(
            "**🏗️ 一、 总体架构**\n正文\n\n1. 普通列表\n\n**二、 环境准备**\n依赖",
            source="proposal.docx",
        )

        # 原因：MinerU 常把普通样式的视觉标题输出为加粗文本，而不是 Markdown #。
        # 作用：保留编号章节结构，但不把所有加粗段落误判为标题。
        self.assertEqual(
            [section.title for section in structure.sections],
            ["一、 总体架构", "二、 环境准备"],
        )
        self.assertIn("1. 普通列表", structure.sections[0].content)

    def test_docx_unstyled_chinese_outline_recovers_substantive_subheadings(self) -> None:
        structure = build_document_structure(
            "腓立比书查经第二十六课\n\n"
            "题目：不抱怨的人，是真正发光的人。\n\n"
            "经文：腓立比书2章14节-16节上半节\n\n"
            "一、破冰话题\n\n"
            "这里有足够长的破冰内容，说明接下来的问题为何值得讨论。\n\n"
            "二、经文解释和讨论\n\n"
            "1.14节，“无论作什么，都不要发怨言、起争论。”\n\n"
            "第一部分有足够长的解释正文，不应和下一部分混在一起。\n\n"
            "2、15节上半节“使你们无可指摘，诚实无伪。”\n\n"
            "第二部分也有足够长的解释正文，用来确认这是一组章节标题。\n\n"
            "三、生活运用（15分钟）\n\n"
            "1. 苹果\n\n"
            "2. 梨\n\n"
            "四、彼此代祷（10分钟）",
            source="lesson.docx",
        )

        self.assertEqual(
            [section.section_path for section in structure.sections],
            [
                ("Document content",),
                ("一、破冰话题",),
                ("二、经文解释和讨论",),
                (
                    "二、经文解释和讨论",
                    "1.14节，“无论作什么，都不要发怨言、起争论。”",
                ),
                (
                    "二、经文解释和讨论",
                    "2、15节上半节“使你们无可指摘，诚实无伪。”",
                ),
                ("三、生活运用（15分钟）",),
                ("四、彼此代祷（10分钟）",),
            ],
        )
        application = next(
            section
            for section in structure.sections
            if section.title == "三、生活运用（15分钟）"
        )
        self.assertIn("1. 苹果", application.content)
        self.assertIn("2. 梨", application.content)

    def test_docx_unstyled_decimal_outline_builds_nested_paths(self) -> None:
        structure = build_document_structure(
            "第一章：背景\n\n"
            "本章正文提供足够长的介绍内容。\n\n"
            "1.1 目标\n\n"
            "目标部分有足够长的正文来证明它不是一个普通列表项。\n\n"
            "1.2 范围\n\n"
            "范围部分也有足够长的正文来形成连续的同级标题。\n\n"
            "第二章：结论\n\n"
            "结论正文。",
            source="requirements.docx",
        )

        self.assertEqual(
            [section.section_path for section in structure.sections],
            [
                ("第一章：背景",),
                ("第一章：背景", "1.1 目标"),
                ("第一章：背景", "1.2 范围"),
                ("第二章：结论",),
            ],
        )

    def test_heading_free_document_gets_one_content_section(self) -> None:
        structure = build_document_structure(
            "A heading-free note.\n\nSecond paragraph.",
            source="notes.md",
        )

        self.assertEqual(len(structure.sections), 1)
        self.assertEqual(structure.sections[0].title, "Document content")
        self.assertIn("Second paragraph", structure.sections[0].content)

    def test_docx_fallback_preserves_heading_styles_as_markdown(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "structured.docx"
            document = Document()
            document.add_heading("Architecture", level=1)
            document.add_paragraph("Overview")
            document.add_heading("Components", level=2)
            document.add_paragraph("Planner and Executor")
            document.save(path)

            with patch(
                "qwopus_agent.documents.parser.parse_document_with_mineru",
                side_effect=MinerUUnavailableError("offline"),
            ):
                parsed = parse_document(path)

        self.assertIn("# Architecture", parsed.markdown)
        self.assertIn("## Components", parsed.markdown)

    def test_chunks_stay_inside_sections_and_respect_token_limit(self) -> None:
        structure = build_document_structure(
            "# First\n\n" + "alpha detail " * 100 + "\n\n# Second\n\n" + "beta fact " * 80,
            source="long.md",
        )

        chunked = chunk_document_structure(structure, max_tokens=64, overlap_tokens=8)

        self.assertGreater(len(chunked.chunks), 2)
        self.assertTrue(all(chunk.token_count <= 64 for chunk in chunked.chunks))
        self.assertTrue(
            all(
                "beta" not in chunk.content
                for chunk in chunked.chunks
                if chunk.section_path == ("First",)
            )
        )
        self.assertTrue(
            all(
                "alpha" not in chunk.content
                for chunk in chunked.chunks
                if chunk.section_path == ("Second",)
            )
        )

    def test_document_store_persists_original_markdown_structure_and_chunks(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            original = Path(tmpdir) / "source.md"
            original.write_text("# Chapter\n\nPersistent content.", encoding="utf-8")
            structure = chunk_document_structure(
                build_document_structure(original.read_text(), source=original.name)
            )
            store = DocumentStore(Path(tmpdir) / "documents")

            directory = store.persist(
                original_path=original,
                markdown=original.read_text(encoding="utf-8"),
                structure=structure,
                metadata={"parser": "markdown"},
            )
            reloaded = store.load_structure(structure.document_id)

            self.assertEqual(reloaded, structure)
            self.assertTrue((directory / "original.md").exists())
            self.assertTrue((directory / "normalized.md").exists())
            self.assertTrue((directory / "chunks.jsonl").read_text(encoding="utf-8"))
            inventory = store.list_documents()
            self.assertEqual(len(inventory), 1)
            self.assertEqual(inventory[0].source, original.name)
            self.assertEqual(inventory[0].section_count, len(structure.sections))

    def test_hierarchical_summary_covers_first_and_last_chapters(self) -> None:
        structure = chunk_document_structure(
            build_document_structure(
                "# Opening\n\nThe project begins with discovery.\n\n"
                "# Closing\n\nThe final decision is to launch in September.",
                source="plan.md",
            ),
            max_tokens=32,
            overlap_tokens=4,
        )

        summary = summarize_document(structure, map_tokens=24, document_tokens=120)

        self.assertIn("project begins", summary.document_summary)
        self.assertIn("final decision", summary.document_summary)
        self.assertTrue(all(section.chunk_ids for section in summary.section_summaries))


if __name__ == "__main__":
    unittest.main()
